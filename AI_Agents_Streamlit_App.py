import streamlit as st
from crewai import Agent, Task, Crew, Process
import os
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
import base64

load_dotenv()

os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")


def generate_docx(result):
    doc = Document()
    doc.add_heading('Health Care Diagonsis and Treatment Recommendations', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def get_download_link(bio, filename):
    """
    Generate a downloadable link for a binary file (e.g., a .docx document).
    """
    # Encode binary content directly into base64
    b64 = base64.b64encode(bio.read()).decode('utf-8')
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'


st.set_page_config(layout="wide")
st.title('AI Agents to empower agents')

gender = st.selectbox('Select Gender', ('Male', 'Female', 'Other'))
age = st.number_input('Enter age', min_value=0, max_value=120, value=25)
symptoms = st.text_area('Enter Symptoms', 'e.g.,fever,cough,headache')
medical_history = st.text_area('Enter medical History', 'e.g,diabetes,hypertension')

search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

llm = ChatOpenAI(
    model='gpt-3.5-turbo-16k',
    temperature=0.1,
    max_tokens=8000
)

# define agents
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis",
    backstory='This agent specializes in diagnosing medical conditions based on patient reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues ',
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
    backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
    verbose=True,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)
diagnose_task = Task(
    description=(
        "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
        "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
        "3. Limit the diagnosis to the most likely conditions."
    ),
    expected_output="A preliminary diagnosis with a list of possible conditions",
    agent=diagnostician
)

treatment_task = Task(
    description=(
        "1. Based on the diagnosis, recommend appropriate treatment plans step by step.\n"
        "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
        "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
    ),
    expected_output="A comprehensive treatment plan tailored to the patients needs",
    agent=treatment_advisor
)

crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=True
)

if st.button("Get Diagnosis and Treatment plan"):
    with st.spinner('Generating recommendations'):
        result = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})

        # Convert result to string if it is not already
        result_str = str(result)

        st.write(result_str)

        # Pass the string result to generate_docx
        docx_file = generate_docx(result_str)

        download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
        st.markdown(download_link, unsafe_allow_html=True)
